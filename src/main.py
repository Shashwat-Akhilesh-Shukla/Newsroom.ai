"""
Main entry point for AI Newsroom.

Run the complete multi-agent newsroom workflow.
"""

import logging
import sys
import argparse
import asyncio
import json
import shutil
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
sys.stdout.reconfigure(encoding='utf-8')

# Load environment variables explicitly
load_dotenv()

# Add project root to path so 'src.' imports work
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.graph import run_newsroom, stream_newsroom
from src.state import create_initial_state
from src.utils.config import get_config
from src.utils.events import RedisLogHandler
from src.observability import setup_tracing
from src.storage.memory import SystemMemory

run_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
log_file_name = f'newsroom_{run_timestamp}.log'

# Configure logging
redis_handler = RedisLogHandler()
redis_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(log_file_name, encoding='utf-8'),
        redis_handler
    ]
)

logger = logging.getLogger(__name__)


async def main():
    """Main execution function."""
    parser = argparse.ArgumentParser(description='AI Newsroom - Multi-Agent Content Creation')
    parser.add_argument('--stream', action='store_true', help='Stream workflow execution')
    parser.add_argument('--debug', action='store_true', help='Enable debug logging')
    parser.add_argument('--output', type=str, help='Output file for final article')
    
    args = parser.parse_args()
    
    # Set debug logging if requested
    if args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Check configuration
    config = get_config()
    if not config.llm.api_key:
        logger.error("❌ No LLM API key found!")
        logger.error("Set GEMINI_API_KEY in your environment")
        sys.exit(1)

    # Initialise LangSmith tracing (no-op if LANGCHAIN_API_KEY not set)
    tracing_active = setup_tracing()
    if not tracing_active:
        logger.info(
            "ℹ️  LangSmith tracing disabled. "
            "Set LANGCHAIN_TRACING_V2=true + LANGCHAIN_API_KEY to enable."
        )
    
    logger.info("=" * 70)
    logger.info("AI NEWSROOM - Multi-Agent Content Creation System")
    logger.info("=" * 70)
    logger.info(f"LLM Provider: {config.llm.provider}")
    logger.info(f"Model: {config.llm.model}")
    logger.info(f"Scout Confidence Threshold: {config.agents.scout_confidence_threshold}")
    logger.info(f"Max Revision Loops: {config.agents.max_revision_loops}")
    logger.info("=" * 70)
    
    # Create initial state
    initial_state = create_initial_state()
    
    try:
        if args.stream:
            # Stream mode - show progress
            logger.info("\n🚀 Starting workflow (streaming mode)...\n")
            
            final_state = None
            async for state_update in stream_newsroom(initial_state):
                # Log state updates
                for node_name, node_state in state_update.items():
                    logger.info(f"📍 {node_name.upper()}: {node_state.get('workflow_stage', 'processing')}")
                final_state = node_state
            
        else:
            # Standard mode - run to completion
            logger.info("\n🚀 Starting workflow...\n")
            final_state = await run_newsroom(initial_state)
        
        # Display results
        logger.info("\n" + "=" * 70)
        logger.info("WORKFLOW COMPLETE")
        logger.info("=" * 70)
        
        # Determine output directory
        output_dir = Path("output") / f"run_{run_timestamp}"
        if args.output:
            out_p = Path(args.output)
            output_dir = out_p if out_p.is_dir() else out_p.parent / f"run_{run_timestamp}"
            
        save_artifacts(final_state, output_dir, log_file_name)
        
        if final_state.get("publish_ready"):
            logger.info("✅ Status: PUBLISHED")
            logger.info(f"📝 Topic: {final_state.get('topic')}")
            logger.info(f"📊 Word Count: {len(final_state.get('draft', '').split())}")
            logger.info(f"🔄 Draft Versions: {final_state.get('draft_version', 0)}")
            logger.info(f"📚 Research Sources: {len(final_state.get('research_notes', []))}")
                
            # Log successful memory
            try:
                memory = SystemMemory()
                memory.mark_topic_published(
                    final_state.get('topic', 'Unknown'),
                    metadata={"word_count": len(final_state.get('draft', '').split())}
                )
            except Exception as e:
                logger.error(f"Failed to record memory publication: {e}")
        
        else:
            logger.info("❌ Status: NOT PUBLISHED")
            logger.info(f"⚠️ Reason: Workflow ended at {final_state.get('workflow_stage', 'unknown')}")
            
            # Show feedback if available
            if final_state.get("critic_feedback"):
                logger.info("\n📋 Skeptic Feedback:")
                for feedback in final_state["critic_feedback"][-2:]:
                    logger.info(f"   {feedback}")
            
            if final_state.get("editor_comments"):
                logger.info("\n📋 Editor Comments:")
                for comment in final_state["editor_comments"][-2:]:
                    logger.info(f"   {comment}")
            
            # Log rejection memory
            try:
                memory = SystemMemory()
                topic = final_state.get('topic')
                if topic:
                    reason = "Workflow failed."
                    if final_state.get("critic_feedback"):
                        reason = str(final_state["critic_feedback"][-1])
                    elif final_state.get("editor_comments"):
                        reason = str(final_state["editor_comments"][-1])
                    memory.mark_topic_rejected(topic, reason)
            except Exception as e:
                logger.error(f"Failed to record memory rejection: {e}")
        
        logger.info("=" * 70)
        
        # Exit with appropriate code
        sys.exit(0 if final_state.get("publish_ready") else 1)
        
    except KeyboardInterrupt:
        logger.info("\n\n⚠️ Workflow interrupted by user")
        sys.exit(130)
        
    except Exception as e:
        logger.error(f"\n\n❌ Workflow failed: {e}", exc_info=True)
        sys.exit(1)


def save_artifacts(state: dict, base_dir: Path, log_file: str):
    """
    Save all requested artifacts to the specified directory.
    """
    try:
        base_dir.mkdir(parents=True, exist_ok=True)
        
        # 1. article.md
        article_content = []
        article_content.append("---")
        article_content.append(f"title: {state.get('topic', 'Untitled')}")
        
        publishing_metadata = state.get("publishing_metadata", {})
        if publishing_metadata:
            article_content.append(f"author: {publishing_metadata.get('author', 'AI Newsroom')}")
            article_content.append(f"date: {publishing_metadata.get('published_date', '')}")
            article_content.append(f"keywords: {', '.join(publishing_metadata.get('keywords', []))}")
            article_content.append(f"reading_time: {publishing_metadata.get('reading_time_minutes', 0)} min")
        
        article_content.append("---")
        article_content.append("")
        
        article_content.append(state.get("draft", ""))
        
        if state.get("research_notes"):
            article_content.append("\n\n---\n")
            article_content.append("## References\n")
            
            for i, note in enumerate(state["research_notes"][:10], 1):
                citation = note.get("citation", "")
                url = note.get("url", "")
                if citation:
                    article_content.append(f"{i}. {citation}")
                    if url:
                        article_content.append(f"   {url}")
                    article_content.append("")
                    
        md_text = '\n'.join(article_content)
        (base_dir / "article.md").write_text(md_text, encoding='utf-8')
        
        # 2. article.docx
        try:
            import docx
            
            doc = docx.Document()
            doc.add_heading(state.get('topic', 'Untitled'), 0)
            
            for paragraph in state.get('draft', '').split('\n'):
                if paragraph.strip():
                    if paragraph.startswith('#'):
                        level = min(len(paragraph) - len(paragraph.lstrip('#')), 4)
                        doc.add_heading(paragraph.lstrip('#').strip(), level)
                    else:
                        doc.add_paragraph(paragraph)
            
            if state.get("research_notes"):
                doc.add_heading("References", 1)
                for note in state["research_notes"][:10]:
                    if note.get("citation"):
                        p = doc.add_paragraph(note["citation"], style='List Bullet')
                        if note.get("url"):
                            p.add_run(f" ({note['url']})")
                            
            doc.save(base_dir / "article.docx")
        except ImportError:
            logger.warning("python-docx not installed, skipping article.docx generation.")
            
        # 3. research_notes.json
        (base_dir / "research_notes.json").write_text(
            json.dumps(state.get("research_notes", []), indent=2), encoding="utf-8"
        )
        
        # 4. run_report.json
        report = {
            "topic": state.get("topic"),
            "status": "PUBLISHED" if state.get("publish_ready") else "NOT PUBLISHED",
            "workflow_stage": state.get("workflow_stage"),
            "draft_version": state.get("draft_version"),
            "revision_count": state.get("revision_count"),
            "iteration_counts": state.get("iteration_counts", {}),
            "created_at": state.get("created_at"),
            "updated_at": state.get("updated_at"),
            "metadata": state.get("metadata", {})
        }
        (base_dir / "run_report.json").write_text(
            json.dumps(report, indent=2, default=str), encoding="utf-8"
        )
        
        # 5. logs.txt
        if Path(log_file).exists():
            shutil.copy2(log_file, base_dir / "logs.txt")
            
        logger.info(f"💾 Artifacts stored in: {base_dir}")
        
    except Exception as e:
        logger.error(f"Failed to save artifacts: {e}", exc_info=True)


def run():
    """Sync entry point for the console script (pyproject.toml)."""
    asyncio.run(main())


if __name__ == "__main__":
    run()
